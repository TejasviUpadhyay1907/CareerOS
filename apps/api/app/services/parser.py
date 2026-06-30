"""Resume parser service for text extraction."""
import io
import re
from typing import Optional

from app.core.logging import get_logger

logger = get_logger(__name__)


class ParserService:
    """Service for parsing resume files and extracting text."""

    async def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text
        """
        try:
            import PyPDF2

            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            logger.info(f"Extracted {len(text)} characters from PDF")
            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise

    async def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file including tables."""
        try:
            import docx

            doc = docx.Document(io.BytesIO(file_content))
            parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Extract tables (many resumes put skills/projects in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            text = "\n".join(parts)
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise

    def clean_text(self, text: str) -> str:
        """Clean extracted text while preserving structure."""
        # Remove control characters but keep newlines and tabs
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        # Reduce multiple spaces to single space (but keep newlines)
        text = re.sub(r"[ \t]+", " ", text)
        # Reduce 3+ newlines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def detect_sections(self, text: str) -> dict[str, str]:
        """Detect sections in resume text.

        Args:
            text: Resume text

        Returns:
            Dictionary of section names and their content
        """
        sections = {}
        current_section = "summary"
        current_content = []

        # Common section headers
        section_patterns = {
            "experience": r"(?i)^(experience|work experience|professional experience|employment history)",
            "education": r"(?i)^(education|academic background|academic)",
            "skills": r"(?i)^(skills|technical skills|core competencies|technologies)",
            "projects": r"(?i)^(projects|portfolio|personal projects)",
            "certifications": r"(?i)^(certifications|certificates|credentials)",
            "languages": r"(?i)^(languages|language proficiency)",
            "achievements": r"(?i)^(achievements|accomplishments|awards)",
            "links": r"(?i)^(links|social|contact|github|linkedin)",
        }

        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line is a section header
            matched_section = None
            for section, pattern in section_patterns.items():
                if re.match(pattern, line):
                    matched_section = section
                    break

            if matched_section:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = matched_section
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content)

        logger.info(f"Detected {len(sections)} sections")
        return sections


# Global instance
parser_service = ParserService()
